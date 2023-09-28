# -*- coding: utf-8 -*-
import random,os,json,io,re,zipfile,tempfile
import pandas as pd
import streamlit as st
import streamlit_toggle as tog
from langchain import HuggingFaceHub
from langchain.llms import OpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
import openai 
import fitz
import docx
from gtts import gTTS
import PyPDF2
from PyPDF2 import PdfReader
from utils import text_to_docs
from langchain import PromptTemplate, LLMChain
from langchain.chat_models import ChatOpenAI
from langchain.chains import ConversationChain
from langchain.memory import ConversationBufferMemory
from langchain.memory import ConversationSummaryBufferMemory
from langchain.chains.conversation.prompt import ENTITY_MEMORY_CONVERSATION_TEMPLATE
from langchain.chains.conversation.memory import ConversationEntityMemory
from langchain.callbacks import get_openai_callback
from io import StringIO
from io import BytesIO
from usellm import Message, Options, UseLLM
from huggingface_hub import login
#from playsound import playsound
#from langchain.text_splitter import CharacterTextSplitter
#from langchain.embeddings.openai import OpenAIEmbeddings
#from langchain.chains.summarize import load_summarize_chain
#import os
#import pyaudio
#import wave
#from langchain.document_loaders import UnstructuredPDFLoader
#import streamlit.components.v1 as components
#from st_custom_components import st_audiorec, text_to_docs
#import sounddevice as sd
#from scipy.io.wavfile import write
                                
# Setting Env
# if st.secrets["OPENAI_API_KEY"] is not None:
#     os.environ["OPENAI_API_KEY"] = st.secrets["OPENAI_API_KEY"]
# else:
#     os.environ["OPENAI_API_KEY"] = os.environ.get("OPENAI_API_KEY")
os.environ["OPENAI_API_KEY"] = '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'

@st.cache_data
def show_pdf(file_path):
    with open(file_path,"rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
        pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="1000px" type="application/pdf"></iframe>'
        st.markdown(pdf_display, unsafe_allow_html=True)

@st.cache_data
def pdf_to_bytes(pdf_file_):
    with open(pdf_file_,"rb") as pdf_file:
        pdf_content = pdf_file.read()
        pdf_bytes_io = io.BytesIO(pdf_content)
    return pdf_bytes_io

@st.cache_data
def read_pdf_files(path):
    pdf_files =[]
    directoty_path = path
    files = os.listdir(directoty_path)
    for file in files:
            pdf_files.append(file)
    return pdf_files


@st.cache_data
def merge_pdfs(pdf_list):
    """
    Helper function to merge PDFs
    """
    pdf_merger = PyPDF2.PdfMerger()
    for pdf in pdf_list:
        pdf_document = PyPDF2.PdfReader(pdf)
        pdf_merger.append(pdf_document)
    output_pdf = BytesIO()
    pdf_merger.write(output_pdf)
    pdf_merger.close()
  
    return output_pdf


@st.cache_data
def usellm(prompt):
    """
    Getting GPT-3.5 Model into action
    """
    service = UseLLM(service_url="https://usellm.org/api/llm")
    messages = [
      Message(role="system", content="You are a financial analyst, who is an expert at giving responses based on the context."),
      Message(role="user", content=f"{prompt}"),
      ]
    options = Options(messages=messages)
    response = service.chat(options)
    return response.content

# # Setting Config for Llama-2
# login(token=st.secrets["HUGGINGFACEHUB_API_TOKEN"])
# os.environ["HUGGINGFACEHUB_API_TOKEN"] = st.secrets["HUGGINGFACEHUB_API_TOKEN"]

login(token="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")
os.environ["HUGGINGFACEHUB_API_TOKEN"]="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"

llama_13b = HuggingFaceHub(
            repo_id="meta-llama/Llama-2-13b-chat-hf",
            model_kwargs={"temperature":0.01, 
                        "min_new_tokens":100, 
                        "max_new_tokens":500})

memory = ConversationSummaryBufferMemory(llm= llama_13b, max_token_limit=500)
conversation = ConversationChain(llm= llama_13b, memory=memory,verbose=False)


@st.cache_data
def llama_llm(_llm,prompt):

    response = _llm.predict(prompt)
    return response

@st.cache_data
def process_text(text):
    # Add your custom text processing logic here
    processed_text = text
    return processed_text

@st.cache_resource
def embed(model_name):
    hf_embeddings = HuggingFaceEmbeddings(model_name=model_name)
    return hf_embeddings

@st.cache_data
def embedding_store(pdf_files):
    pdf_only =[]
    text = ""
    for file in pdf_files:
      if file.endswith('.pdf'):
        pdf_only.append(file)       
      
    merged_pdf = merge_pdfs(pdf_only)
    final_pdf = PyPDF2.PdfReader(merged_pdf)
    for page in final_pdf.pages:
        text += page.extract_text()
      
    for file in pdf_files:
      if file.endswith('xlsx'):
        df = pd.read_excel(file, engine='openpyxl')
        # Find the row index where the table data starts
        data_start_row = 0  # Initialize to 0
        for i, row in df.iterrows():
            if row.notna().all():
                data_start_row = i
                break
              
        if data_start_row>0:  
            df.columns = df.iloc[data_start_row]
          
        
        # Extract the text content above the data
        text += "\n".join(df.iloc[:data_start_row].apply(lambda x: "\t".join(map(str, x)), axis=1)).replace('nan','')
        
        df1 = df.iloc[data_start_row+1:]
        text_buffer = StringIO()
        df1.to_csv(text_buffer, sep='\t', index=False)
        text += "\n\n"+ text_buffer.getvalue()
        text_buffer.close()
        
    texts =  text_splitter.split_text(text)
    docs = text_to_docs(texts)
    docsearch = FAISS.from_documents(docs, hf_embeddings)
    return docs, docsearch
    
@st.cache_data
def merge_and_extract_text(pdf_list):
    """
    Helper function to merge PDFs and extract text
    """
    pdf_merger = PyPDF2.PdfMerger()
    for pdf in pdf_list:
        with open(pdf, 'rb') as file:
            pdf_merger.append(file)
    output_pdf = BytesIO()
    pdf_merger.write(output_pdf)
    pdf_merger.close()
    
    # Extract text from merged PDF
    merged_pdf = PyPDF2.PdfReader(output_pdf)
    all_text = []
    for page in merged_pdf.pages:
        text = page.extract_text()
        all_text.append(text)
    
    return ' '.join(all_text)

def reset_session_state():
    session_state = st.session_state
    session_state.clear()

# def merge_and_extract_text(pdf_list):
#     merged_pdf = fitz.open()
#     # Merge the PDF files
#     for pdf_file in pdf_list:
#         pdf_document = fitz.open(pdf_file)
#         merged_pdf.insert_pdf(pdf_document)
#     # Create an empty string to store the extracted text
#     merged_text = ""
#     # Extract text from each page of the merged PDF
#     for page_num in range(merged_pdf.page_count):
#         page = merged_pdf[page_num]
#         text = page.get_text()
#         merged_text += text
#     # Close the merged PDF
#     merged_pdf.close()
#     return merged_text


@st.cache_data
def render_pdf_as_images(pdf_file):
    """
    Helper function to render PDF pages as images
    """
    pdf_images = []
    pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        img = page.get_pixmap()
        img_bytes = img.tobytes()
        pdf_images.append(img_bytes)
    pdf_document.close()
    return pdf_images

# To check if pdf is searchable
def is_searchable_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            if page.extract_text():
                return True

    return False


# Function to add checkboxes to the DataFrame
@st.cache_data
def add_checkboxes_to_dataframe(df):
    # Create a new column 'Select' with checkboxes
    checkbox_values = [True] * (len(df) - 1) + [False]  # All True except the last row
    df['Select'] = checkbox_values
    return df

# convert scanned pdf to searchable pdf
def convert_scanned_pdf_to_searchable_pdf(input_file):
    """
     Convert a Scanned PDF to Searchable PDF

    """
    # Convert PDF to images
    print("Running OCR")
    images = convert_from_path(input_file)

    # Preprocess images using OpenCV
    for i, image in enumerate(images):
        # Convert image to grayscale
        image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2GRAY)

        # Apply thresholding to remove noise
        _, image = cv2.threshold(image, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

        # Enhance contrast
        image = cv2.equalizeHist(image)

        # Save preprocessed image
        cv2.imwrite(f'{i}.png', image)

    # Perform OCR on preprocessed images using Tesseract
    text = ''
    for i in range(len(images)):
        image = cv2.imread(f'{i}.png')
        text += pytesseract.image_to_string(image)
    
    return text


# Setting globals
if "visibility" not in st.session_state:
    st.session_state.visibility = "visible"
    st.session_state.disabled = True
if "stored_session" not in st.session_state:
    st.session_state["stored_session"] = []
if "tmp_table_gpt_fd" not in st.session_state:
    st.session_state.tmp_table_gpt_fd=pd.DataFrame()
if "tmp_table_llama_fd" not in st.session_state:
    st.session_state.tmp_table_llama_fd=pd.DataFrame()
if "tmp_table_gpt_aml" not in st.session_state:
    st.session_state.tmp_table_gpt_aml=pd.DataFrame()
if "tmp_table_llama_aml" not in st.session_state:
    st.session_state.tmp_table_llama_aml=pd.DataFrame()

if "tmp_narr_table_gpt" not in st.session_state:
    st.session_state.tmp_narr_table_gpt=pd.DataFrame()
if "tmp_narr_table_llama" not in st.session_state:
    st.session_state.tmp_narr_table_llama=pd.DataFrame()
  
if "tmp_summary_gpt_fd" not in st.session_state:
    st.session_state["tmp_summary_gpt_fd"] = ''
if "tmp_summary_llama_fd" not in st.session_state:
    st.session_state["tmp_summary_llama_fd"] = ''
if "tmp_summary_gpt_aml" not in st.session_state:
    st.session_state["tmp_summary_gpt_aml"] = ''
if "tmp_summary_llama_aml" not in st.session_state:
    st.session_state["tmp_summary_llama_aml"] = ''

if "tmp_narrative_gpt" not in st.session_state:
    st.session_state["tmp_narrative_gpt"] = ''
if "tmp_narrative_llama" not in st.session_state:
    st.session_state["tmp_narrative_llama"] = ''

if "case_num" not in st.session_state:
    st.session_state.case_num = ''
if "fin_opt" not in st.session_state:
    st.session_state.fin_opt = ''
if "context_1" not in st.session_state:
    st.session_state.context_1 = ''
if "llm" not in st.session_state:
    st.session_state.llm = 'Open-AI'
if "pdf_files" not in st.session_state:
    st.session_state.pdf_files = []


# Apply CSS styling to resize the buttons
st.markdown("""
    <style>
        .stButton button {
            width: 145px;
            height: 35px;
        }
    </style>
""", unsafe_allow_html=True)

@st.cache_data
def add_footer_with_fixed_text(doc, footer_text):
    # Create a footer object
    footer = doc.sections[0].footer

    # Add a paragraph to the footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    # Set the fixed text in the footer
    paragraph.text = footer_text

    # Add a page number field to the footer
    run = paragraph.add_run()
    fld_xml = f'<w:fldSimple {nsdecls("w")} w:instr="PAGE"/>'
    fld_simple = parse_xml(fld_xml)
    run._r.append(fld_simple)

    # Set the alignment of the footer text
    paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

@st.cache_data
def create_filled_box_with_text(color, text):
    box_html = f'<div style="flex: 1; height: 100px; background-color: {color}; display: flex; align-items: center; justify-content: center;">{text}</div>'
    st.markdown(box_html, unsafe_allow_html=True)

@st.cache_data
def create_zip_file(file_paths, zip_file_name):
    with zipfile.ZipFile(zip_file_name, 'w') as zipf:
        for file_path in file_paths:
            zipf.write(file_path, os.path.basename(file_path))

# Addding markdown styles(Global)
st.markdown("""
<style>
.big-font {
    font-size:60px !important;
}
</style>
""", unsafe_allow_html=True)


# Set Sidebar
st.markdown("""
<style>
    [data-testid=stSidebar] {
        background-color: FBFBFB;
    }
</style>
""", unsafe_allow_html=True)

#Adding llm type-> st.session_state.llm
st.session_state.llm = st.radio("",options = pd.Series(["","Open-AI","Open-Source"]), horizontal=True)
st.markdown("---")

st.title("Suspicious Activity Reporting Assistant")
with st.sidebar:
    # st.sidebar.write("This is :blue[test]")
    # Navbar
    st.markdown('<link rel="stylesheet" href="https://maxcdn.bootstrapcdn.com/bootstrap/4.0.0/css/bootstrap.min.css" integrity="sha384-Gn5384xqQ1aoWXA+058RXPxPg6fy4IWvTNh0E263XmFcJlSAwiGgFAW/dAiS6JXm" crossorigin="anonymous">', unsafe_allow_html=True)

    st.markdown("""
    <nav class="navbar fixed-top navbar-expand-lg navbar-dark" style="background-color: #000000;">
    <button class="navbar-toggler" type="button" data-toggle="collapse" data-target="#navbarNav" aria-controls="navbarNav" aria-expanded="false" aria-label="Toggle navigation">
        <span class="navbar-toggler-icon"></span>
    </button>
    <style>
    .navbar-brand img {
      max-height: 50px; /* Adjust the height of the logo */
      width: auto; /* Allow the width to adjust based on the height */
    }
    </style>
    <div class="collapse navbar-collapse" id="navbarNav">
        <ul class="navbar-nav">
        <li class="nav-item active">
            <a class="navbar-brand" href="#">
                <img src="https://www.exlservice.com/themes/exl_service/exl_logo_rgb_orange_pos_94.png" width="50" height="30" alt="">
                <span class="sr-only">(current)</span>
                <strong>| Operations Process Automation</strong>
            </a>
        </li>
        </ul>
    </div>
    </nav>
    """, unsafe_allow_html=True)

    st.markdown("""
    <nav class="navbar fixed-bottom navbar-expand-lg navbar-dark" style="background-color: #000000;">
    <button class="navbar-toggler" type="button" data-toggle="collapse" data-target="#navbarNav" aria-controls="navbarNav" aria-expanded="false" aria-label="Toggle navigation">
        <span class="navbar-toggler-icon"></span>
    </button>
    <div class="collapse navbar-collapse" id="navbarNav">
        <ul class="navbar-nav">
        <li class="nav-item active">
        <!--p style='color: white;'><b>Powered by EXL</b></p--!>
        <p style='color: white;'> <strong>Powered by EXL</strong> </p>
            <!--a class="nav-link disabled" href="#">
                <img src="https://www.exlservice.com/themes/exl_service/exl_logo_rgb_orange_pos_94.png" width="50" height="30" alt="">
                <span class="sr-only">(current)</span>
            </a--!>
        </li>
        </ul>
    </div>
    </nav>
    """, unsafe_allow_html=True)

    # Add the app name
    st.sidebar.markdown('<p class="big-font">SARA</p>', unsafe_allow_html=True)
    # st.sidebar.header("SARA")
    st.markdown("---")

    # Add a drop-down for case type
    options = ["Select Case Type", "Fraud transaction dispute", "AML"]
    selected_option_case_type = st.sidebar.selectbox("", options)
    st.markdown("---")
    
    # Add a single dropdown
    options = ["Select Case ID", "SAR-2023-24680", "SAR-2023-13579", "SAR-2023-97531", "SAR-2023-86420", "SAR-2023-24681"]
    selected_option = st.sidebar.selectbox("", options)
    # Add the image to the sidebar below options
    st.sidebar.image("MicrosoftTeams-image (7).png", use_column_width=True)

    
# Assing action to the main section
if selected_option_case_type == "Select Case Type":
    st.header("")
elif selected_option_case_type == "Fraud transaction dispute":    
    
    st.markdown("### :blue[Fraud transaction dispute]")

    # Redirect to Merge PDFs page when "Merge PDFs" is selected
    if selected_option == "SAR-2023-24680":
        st.session_state.case_num = "SAR-2023-24680"
        # st.header("Merge Documents")
        # st.write("Upload multiple document files and merge them into one doc.")
    
        # Upload PDF files
        # st.subheader("Upload Case Files")
        # st.markdown(f"**Case No: {st.session_state.case_num}**")
        # st.markdown("""
        #     | Case No.                  | Case Type                 | Customer Name             | Case Status             | Open Date              |
        #     | ------------------------  | ------------------------- | ------------------------- | ------------------------|------------------------|
        #     | SAR-2023-24680            | Fraud Transaction Dispute | John Brown                | In Progress             | 12/10/2020             |
        #     """)
    
        col1,col2 = st.columns(2)
        # Row 1
        with col1:
            st.markdown("**Case number&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** SAR-2023-24680")
            st.markdown("**Customer name  :** John Brown")
    
    
        with col2:
            st.markdown("**Case open date&nbsp;&nbsp;&nbsp;&nbsp;:** Feb 02, 2021")
            st.markdown("**Case type&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** Fraud transaction")
    
    
        # Row 2
        with col1:
            st.markdown("**Customer ID&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** 9659754")
    
    
        with col2:
            st.markdown("**Case Status&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** Open")
    
    
        # Evidence uploader/Fetch    
        st.header("Upload Evidence")
        # Showing files record
        # data_source = ["Customer details", "Transaction details", "Fraud dispute details", "Other Files"]
        # data_file_df = pd.DataFrame(data_source, columns=["File Name"])
        # data_file_df = data_file_df.reset_index(drop=True)
    
        # df_s_with_checkboxes = add_checkboxes_to_dataframe(data_file_df)
    
        # # Iterate through each row and add checkboxes
        # for index, row in df_s_with_checkboxes.iterrows():
        #     if index < -1:
        #         checkbox_state = st.checkbox(f" {row['File Name']}", value=True)
        #         df_s_with_checkboxes.loc[index, 'Select'] = checkbox_state
        #     else:
        #         st.checkbox(f"{row['File Name']}", value=False)
    
    
        if selected_option:
            # Create two columns
            # reading files from local directory from fetch evidence button
            directoty_path = "data/"
            fetched_files = read_pdf_files(directoty_path)

            col1_up, col2_up = st.tabs(["Fetch Evidence", "Upload Evidence"])
            with col1_up:
                # Set the color
                # st.markdown(
                #     """
                #     <div style="display: flex; justify-content: center; align-items: center; height: 48px; border: 1px solid #ccc; border-radius: 5px; background-color: #f2f2f2;">
                #         <span style="font-size: 16px;  ">Fetch Evidence</span>
                #     </div>
                #     """,
                #     unsafe_allow_html=True
                # )
                if 'clicked' not in st.session_state:
                    st.session_state.clicked = False
                
                def set_clicked():
                    st.session_state.clicked = True
                    st.session_state.disabled = True
                
                st.button('Fetch Evidence', on_click=set_clicked)
    
                if st.session_state.clicked:
                    
                   #select box to select file
                    selected_file_name = st.selectbox(":blue[Select a file to View]",fetched_files)
                    st.write("Selected File: ", selected_file_name)
                    st.session_state.disabled = False
                    file_ext = tuple("pdf")
                    if selected_file_name.endswith(file_ext):
                        selected_file_path = os.path.join(directoty_path, selected_file_name)
                        #converting pdf data to bytes so that render_pdf_as_images could read it
                        file = pdf_to_bytes(selected_file_path)
                        pdf_images = render_pdf_as_images(file)
                        #showing content of the pdf
                        st.subheader(f"Contents of {selected_file_name}")
                        for img_bytes in pdf_images:
                            st.image(img_bytes, use_column_width=True)
                    else:
                        selected_file_path = os.path.join(directoty_path, selected_file_name)
                        # This is showing png,jpeg files
                        st.image(selected_file_path, use_column_width=True)
    
    
    
            with col2_up:
                pdf_files = st.file_uploader("", type=["pdf","png","jpeg","docx","xlsx"], accept_multiple_files=True)
                st.session_state.pdf_files = pdf_files
                # showing files
                for uploaded_file in pdf_files:
                    #This code is to show pdf files
                    file_ext = tuple("pdf")
                    if uploaded_file.name.endswith(file_ext):
                        # Show uploaded files in a dropdown
                        # if pdf_files:
                        st.subheader("Uploaded Files")
                        file_names = [file.name for file in pdf_files]
                        selected_file = st.selectbox(":blue[Select a file]", file_names)
                        # Enabling the button
                        st.session_state.disabled = False
                        # Display selected PDF contents
                        if selected_file:
                            selected_pdf = [pdf for pdf in pdf_files if pdf.name == selected_file][0]
                            pdf_images = render_pdf_as_images(selected_pdf)
                            st.subheader(f"Contents of {selected_file}")
                            for img_bytes in pdf_images:
                                st.image(img_bytes, use_column_width=True)
    
                    else:
                        # This is showing png,jpeg files
                        st.image(uploaded_file, use_column_width=True)
    
    #creating temp directory to have all the files at one place for accessing
        tmp_dir_ = tempfile.mkdtemp()
        temp_file_path= []
    
        for uploaded_file in pdf_files:
            file_ext = tuple("pdf")
            if uploaded_file.name.endswith(file_ext):
                file_pth = os.path.join(tmp_dir_, uploaded_file.name)
                with open(file_pth, "wb") as file_opn:
                    file_opn.write(uploaded_file.getbuffer())
                    temp_file_path.append(file_pth)
            else:
                pass
    
    
        for fetched_pdf in fetched_files:
            file_ext = tuple("pdf")
            if fetched_pdf.endswith(file_ext):
                file_pth = os.path.join('data/', fetched_pdf)
                # st.write(file_pth)
                temp_file_path.append(file_pth) 
            else:
                pass
    
        #combining files in fetch evidence and upload evidence
        pdf_files_ = []
        if temp_file_path:
            if pdf_files and fetched_files:
                file_names = [file.name for file in pdf_files]
                file_names = file_names + fetched_files
                pdf_files_ = file_names
            elif fetched_files:
                pdf_files_ = fetched_files
            elif pdf_files:
                pdf_files_ = pdf_files
            else: pass
    
         #This is the embedding model
        model_name = "thenlper/gte-small"
        # model_name = "hkunlp/instructor-large"
        
        # Memory setup for gpt-3.5
        llm = ChatOpenAI(temperature=0.1, model ='gpt-3.5-turbo')
        memory = ConversationSummaryBufferMemory(llm=llm, max_token_limit=500)
        conversation = ConversationChain(llm=llm, memory =memory,verbose=False)

        
        # Adding condition on embedding
        try:
            if temp_file_path:
                hf_embeddings = embed(model_name) 
            else:
                pass
        except NameError:
            pass
        
        # Chunking with overlap
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size = 1000,
            chunk_overlap  = 100,
            length_function = len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        
        
    
    else:
        # Disabling the button
        st.session_state.disabled = True
        st.session_state.case_num = selected_option    
    
    # Creating header
    col1,col2 = st.columns(2)
    with col1:
        st.subheader('Pre-Set Questionnaire')
        # Create a Pandas DataFrame with your data
    
        data = {'Questions': [" What is the victim's name?","What is the suspect's name?",' List the merchant name',' How was the bank notified?',' When was the bank notified?',' What is the fraud type?',' When did the fraud occur?',' Was the disputed amount greater than 5000 USD?',' What type of cards are involved?',' Was the police report filed?']}
        df_fixed = pd.DataFrame(data)
        df_fixed.index = df_fixed.index +1
    with col2:
        # Create a checkbox to show/hide the table
        cols1, cols2, cols3, cols4 = st.columns([1,1,1,1])
        with cols1:
            show_table = tog.st_toggle_switch(label="", 
                                key="Key1", 
                                default_value=False, 
                                label_after = False, 
                                inactive_color = '#D3D3D3', 
                                active_color="#11567f", 
                                track_color="#29B5E8"
                                )
        # Show the table if the checkbox is ticked
        if show_table:
            
            df_fixed["S.No."] = df_fixed.index
            df_fixed = df_fixed.loc[:,['S.No.','Questions']]
            st.markdown(df_fixed.style.hide(axis="index").to_html(), unsafe_allow_html=True)
    
    with st.spinner('Wait for it...'):
        if st.button("Generate Insights",disabled=st.session_state.disabled):
            if temp_file_path is not None:
                # File handling logic
                _, docsearch = embedding_store(temp_file_path)
                if st.session_state.llm == "Open-AI":
                    queries ="Please provide the following information regarding the possible fraud case: What is the name of the customer name,\
                    has any suspect been reported, list the merchant name, how was the bank notified, when was the bank notified, what is the fraud type,\
                    when did the fraud occur, was the disputed amount greater than 5000 USD, what type of cards are involved, was the police report filed,\
                    and based on the evidence, is this a suspicious activity(Summarize all the questions asked prior to this in a detailed manner),that's the answer of\
                    whether this is a suspicious activity\
                    "
                    contexts = docsearch.similarity_search(queries, k=5) 
                    prompts = f" Give a the answer to the below questions as truthfully and in as detailed in the form of sentences\
                    as possible as per given context only,\n\n\
                            What is the victim's name?\n\
                            What is the suspect's name?\n\
                            List the merchant name\n\
                            How was the bank notified?\n\
                            When was the bank notified?\n\
                            What is the fraud type?\n\
                            When did the fraud occur?\n\
                            Was the disputed amount greater than 5000 USD?\n\
                            What type of cards are involved?\n\
                            Was the police report filed?\n\
                        Context: {contexts}\n\
                        Response (in the python dictionary format\
                        where the dictionary key would carry the questions and its value would have a descriptive answer to the questions asked): "
                        
                    response = usellm(prompts)
                    
                   
                    # st.write(response)
                    resp_dict_obj = json.loads(response)
                  
                    query = "Is it a SAR Case?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f'''You need to act as a Financial analyst to check if this is a SAR or not, given the context. You can use the chat_history to get the context.\n\ \
                            A SAR case is usually determined by the following factors: \
                            1. If there is a potential suspect name present. \
                            2. If there is a police report filed. \
                            3. If there is a any fradulent transaction made by the suspect. \
                            Give your answer as yes or no.\n\n\
                           Context: {context_1}\n\
                           Response: '''

                    response = usellm(prompt_1)
                    is_sar = response
                    resp_dict_obj[query] = response
            
                    res_df_gpt = pd.DataFrame(resp_dict_obj.items(), columns=['Question','Answer'])            
                    res_df_gpt_new = res_df_gpt.iloc[:-1]
              
                    try:
                        res_df_gpt_new.reset_index(drop=True, inplace=True)
                        index_ = pd.Series([1,2,3,4,5,6,7,8,9,10])
                        res_df_gpt_new = res_df_gpt_new.set_index([index_])
                        # st.write(res_df_gpt)
              
                    except IndexError: 
                        pass
                    
                    st.table(res_df_gpt_new)
                    # st.write(res_df_gpt)

                    #data stored in  session state
                    st.session_state["tmp_table_gpt_fd"] = pd.concat([st.session_state.tmp_table_gpt_fd, res_df_gpt], ignore_index=True)
                    st.session_state["tmp_narr_table_gpt"] = pd.concat([st.session_state.tmp_narr_table_gpt, res_df_gpt], ignore_index=True)
                
                
                elif st.session_state.llm == "Open-Source":
    
                    chat_history = {}
    
                    query = "What is the victim's name?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 = f'''You are a professional fraud analyst. Perform Name Enitity Recognition to identify the victim's name as accurately as possible, given the context. The victim can also be referenced as the customer with whom the Fraud has taken place.
                    victim's name is the Name provided in Cardholder Information.\n\n\
                            Question: {query}\n\
                            Context: {context_1}\n\
                            Response: (Give me response in one sentence. Do not give me any Explanation or Note)'''
                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response
    
    
                    query = "What is the suspect's name?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f'''You are a professional fraud analyst. You need to check the document and compare if any name discrepencies are present that points towards the suspect who used the card without the consent of the cardholder.
                Hence, Compare the names present in the context. 
                Reply the name of the person who is basically the suspect.\n\n\
                                Context: {context_1}\n\
                                Response: (Give me a concise response in one sentence.Do not give me any Explanation,Note)'''
                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response
    
                    
                    
                    query = "list the merchant name"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 = f'''Perform Name Enitity Recognition to identify Merchant as accurately as possible, given the context. A merchant is a type of business or organization that accepts payments from the customer account. Give a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Give me a concise response. Do not add any Explanation,Note.)'''
                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response
    
    
                    query = "How was the bank notified?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f'''You need to act as a Financial analyst to identify how was the bank notified of the Supicious or Fraud event with in the given context. The means of communication can be a call, an email or in person. Give a concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Give me a concise response in one sentence. Do not give me any further Explanation, Note )'''
                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response
    
                    
                    query = "When was the bank notified?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f'''You need to act as a Financial analyst to identify when the bank was notified of the Fraud. Look for the disputed date. Given the context, provide a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Give me a concise response in one sentence.Do not add any prefix like 'Response' or 'Based on the document'. Do not add any extra Explanation, Note)'''
                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response
                    
    
    
                    query = "What is the Fraud Type?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f''' You need to act as a Financial analyst to identify the type of fraud or suspicious activity has taken place amd summarize it, within the given context. Also mention the exact fraud code. Give a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Give me response in one sentence. Do not add prefix like 'Response' or 'based on the document. Do not give me any Explanation or Note)'''
                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response
    
    
    
    
                    query = "When did the fraud occur?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f''' You need to act as a Financial analyst to identify the when the did the fraud occur i.e., the Transaction Date. Given the context, provide a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Give me a concise response in one sentence. Do not add prefix like 'based on the document. Do not add any further Explanation or Note.)'''
                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response
    
    
                    query = "Was the disputed amount greater than 5000 usd?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f''' You need to act as a Financial analyst to identify the disputed amount and perform a mathematical calculation to check if the disputed amount is greater than 5000 USD or not, given the context. Give a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Provide a concise Response in a single sentence. Do not write any extra [Explanation, Note, Descricption].)'''
                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response
    
    
                    query = "What type of cards are involved?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f''' You need to act as a Financial analyst to identify the type of card and card network involved, given the context. On a higher level the card can be a Credit Visa, Debit Visa Card.Based on the context give a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Give me a concise response in one sentence.Do not add prefix like: ['based on the document']. Do not add any further Explanation, Note.')'''
                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response
    
    
                    query = "was the police report filed?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f''' You need to act as a Financial analyst to identify if the police was reported of the Fraud activity, given the context. Give a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Provide a concise Response in a single sentence. Do not write any extra [Explanation, Note, Descricption].)'''
                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response

                    query = "Is it a SAR Case?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f'''You need to act as a Financial analyst to check if this is a SAR or not, given the context. You can use the chat_history to get the context.\n\ \
                            A SAR case is usually determined by the following factors: \
                            1. If there is a potential suspect name present. \
                            2. If there is a police report filed. \
                            3. If there is a any fradulent transaction made by the suspect. \
                            Give your answer as yes or no.\n\n\
                           Context: {context_1}\n\
                           Response: '''

                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response

                    res_df_llama = pd.DataFrame(chat_history.items(), columns=['Question','Answer'])            
                    res_df_llama_new = res_df_llama.iloc[:-1,]
    
                    try:
                        res_df_llama_new.reset_index(drop=True, inplace=True)
                        index_ = pd.Series([1,2,3,4,5,6,7,8,9,10])
                        res_df_llama_new = res_df_llama_new.set_index([index_])
                        # st.write(res_df_llama)
          
                    except IndexError: 
                        pass

                    st.table(res_df_llama_new)
                    # st.write(res_df_llama)
          
                    st.session_state["tmp_table_llama_fd"] = pd.concat([st.session_state.tmp_table_llama_fd, res_df_llama], ignore_index=True)                
                    st.session_state["tmp_narr_table_llama"] = pd.concat([st.session_state.tmp_narr_table_llama, res_df_llama], ignore_index=True)
                
      
    st.markdown("---")
    
    # For input box outside of template4
    try:
        if temp_file_path:
            docs, docsearch = embedding_store(temp_file_path)
        else:
            pass
    except Exception:
        pass
    
    
    # Text Input
    st.subheader("Ask Additional Questions")
    query = st.text_input(':blue[Please ask below the additional case questions.]',disabled=st.session_state.disabled)
    text_dict = {}
    @st.cache_data
    def LLM_Response():
        llm_chain = LLMChain(prompt=prompt, llm=llm)
        response = llm_chain.run({"query":query, "context":context})
        return response
    if st.session_state.llm == "Open-AI":
        with st.spinner('Getting you information...'):      
            if query:
                # Text input handling logic
                #st.write("Text Input:")
                #st.write(text_input)
    
                context_1 = docsearch.similarity_search(query, k=5)
                st.session_state.context_1 = context_1
                if query.lower() == "what is the victim's name?":
                    prompt_1 = f'''Perform Name Enitity Recognition to identify the Customer name as accurately as possible, given the context. The Customer can also be referenced as the Victim or the person with whom the Fraud has taken place.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: '''
    
                    
                elif query.lower() == "what is the suspect's name?":
                    prompt_1 = f'''Perform Name Enitity Recognition to identify the Suspect name as accurately as possible, given the context. Suspect is the Person who has committed the fraud with the Customer. Respond saying "The Suspect Name is not Present" if there is no suspect in the given context.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: '''
    
                    
                elif query.lower() == "list the merchant name":
                    prompt_1 = f'''Perform Name Enitity Recognition to identify all the Merchant Organizations as accurately as possible, given the context. A merchant is a type of business or organization that accepts payments from the customer account. Give a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: '''
    
                    
                elif query.lower() == "how was the bank notified?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify how was the bank notified of the Supicious or Fraud event with in the given context. The means of communication can be a call, an email or in person. Give a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: '''
    
                    
                elif query.lower() == "when was the bank notified?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify the when the bank was notified of the Fraud i.e., the disputed date. Given the context, provide a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: '''
    
                    
                elif query.lower() == "what type of fraud is taking place?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify the type of fraud or suspicious activity has taken place amd summarize it, within the given context. Also mention the exact fraud code. Give a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: '''
    
                
                elif query.lower() == "when did the fraud occur?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify the when the did the fraud occur i.e., the Transaction Date. Given the context, provide a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: '''
    
                        
                elif query.lower() == "was the disputed amount greater than 5000 usd?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify the disputed amount and perform a mathematical calculation to check if the disputed amount is greater than 5000 or no, given the context. Give a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: '''
    
                    
                elif query.lower() == "what type of cards are involved?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify the type of card and card's brand involved, given the context. On a higher level the card can be a Credit or Debit Card. VISA, MasterCard or American Express, Citi Group, etc. are the different brands with respect to a Credit card or Debit Card . Give a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: '''
    
                    
                elif query.lower() == "was the police report filed?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify if the police was reported of the Fraud activity, given the context. Give a relevant and concise response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: '''
    
                        
                elif query.lower() == "Is this a valid SAR case?":
                    prompt_1 = f''' You need to act as a Financial analyst to check if this is a SAR or not, given the following context, if the transaction amount is less than 5000 USD we cannot categorize this as SAR (Suspicious activity Report).Give a relevant and concise response. \n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: '''
    
                    
                else:
                    prompt_1 = f'''Act as a financial analyst and give concise answer to below Question as truthfully as possible, with given Context.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\                      
                                Response: '''
    
    
                #prompt = PromptTemplate(template=prompt, input_variables=["query", "context"])
                response = usellm(prompt_1) #LLM_Response()
                text_dict[query] = response
                # resp_dict_obj.update(text_dict)
                st.write(response)
                if response:
                    df = pd.DataFrame(text_dict.items(), columns=['Question','Answer'])
                else:
                    df = pd.DataFrame()
    
                st.session_state["tmp_table_gpt_fd"] = pd.concat([st.session_state.tmp_table_gpt_fd, df], ignore_index=True)
                st.session_state.tmp_table_gpt_fd.drop_duplicates(subset=['Question'])
    
    
    elif st.session_state.llm == "Open-Source":
        with st.spinner('Getting you information...'):      
            if query:
                # Text input handling logic
                #st.write("Text Input:")
                #st.write(text_input)
    
                context_1 = docsearch.similarity_search(query, k=5)
                st.session_state.context_1 = context_1
                if query.lower() == "what is the victim's name?":
                    prompt_1 = f'''Perform Name Enitity Recognition to identify the Customer name as accurately as possible, given the context. The Customer can also be referenced as the Victim or the person with whom the Fraud has taken place.
                                Customer/Victim is cardholder, whose card is used without their consent.
                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Provide a concise Response.) '''
    
                    
                elif query.lower() == "what is the suspect's name?":
                    prompt_1 = f''''Perform Name Enitity Recognition to identify the Suspect name as accurately as possible, given the context. Suspect is the Person who has committed the fraud with the Customer. Respond saying "The Suspect Name is not Present" if there is no suspect in the given context.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Give me response in one sentence.Do not give me any Explanation or Note)'''
    
    
                    
                elif query.lower() == "list the merchant name":
                    prompt_1 = f'''Perform Name Enitity Recognition to identify all the Merchant Organizations as accurately as possible, given the context. A merchant is a type of business or organization that accepts payments from the customer account. Give a relevant and concise response.
                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Provide a concise Response without any extra [Explanation, Note, Descricption] below the Response.)'''
    
                    
                elif query.lower() == "how was the bank notified?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify how was the bank notified of the Supicious or Fraud event with in the given context. The means of communication can be a call, an email or in person. Give a relevant and concise response.
                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response:(Provide a concise Response.) '''
    
                    
                elif query.lower() == "when was the bank notified?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify the when the bank was notified of the Fraud i.e., the disputed date. Given the context, provide a relevant and concise response.
                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Provide a concise Response.)'''
    
                    
                elif query.lower() == "what type of fraud is taking place?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify the type of fraud or suspicious activity has taken place amd summarize it, within the given context. Also mention the exact fraud code. Give a relevant and concise response.
                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Provide a concise Response without any extra [Explanation, Note, Descricption] below the Response.)'''
    
                
                elif query.lower() == "when did the fraud occur?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify the type of card and card network involved, given the context. On a higher level the card can be a Credit Visa, Debit Visa Card.Based on the context give a relevant and concise response..
                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Provide a concise Response without any extra [Explanation, Note, Descricption] below the Response.)'''
    
                        
                elif query.lower() == "was the disputed amount greater than 5000 usd?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify the disputed amount and perform a mathematical calculation to check if the disputed amount is greater than 5000 or no, given the context. Give a relevant and concise response.
                                Kindly do not provide any extra [Explanation, Note, Description] block below the Response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response:(Provide a concise Response without any extra [Explanation, Note, Descricption] below the Response.) '''
    
                    
                elif query.lower() == "what type of cards are involved?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify the type of Card and Card Network involved, given the context. On a higher level the card can be a Dedit, Crebit Card. VISA, MasterCard, American Express, Citi Group, etc. are the different brands with respect to a Credit Card or Debit Card . Give a relevant and concise response.
                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response:(Act like a professional and provide me a concise Response . Do not add any extra [Explanation, Note, Descricption] below the context.) '''
    
                    
                elif query.lower() == "was the police report filed?":
                    prompt_1 = f''' You need to act as a Financial analyst to identify if the police was reported of the Fraud activity, given the context. Give a relevant and concise response.
                                Do not provide any extra [Explanation, Note] block below the Response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Provide a concise Response without any extra [Explanation, Note, Descricption] below the Response.)'''
    
                elif query.lower() == "is this a valid sar case?":
                    prompt_1 =  f''' You are a Fraud Analyst.Check if there is evidence for this case to address as SAR or not. A SAR case is a case of financial Suspicious/Fraud Activity which can be observed given the context.
                                If there is any activity without the consent of the cardholder, also if there is a suspect who used the card without the consent.
                                Then we can address this as a valid SAR case.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response: (Provide a concise response in single sentence.Do not add prefix like ['Respone', 'based on the document']. Do not add any further Explanation,Note.)'''        
                
                
                elif query.lower() == "is there any evidence of a sar case?":
                    prompt_1 = f''' You are a Fraud Analyst.Check if there is evidence for this case to address as SAR or not. A SAR case is a case of financial Suspicious/Fraud Activity which can be observed given the context.
                                If there is any activity without the consent of the cardholder, also if there is a suspect who used the card without the consent.
                                Then we can address this as a SAR case.Give a concise response with the suspect name. \n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\
                                Response:(Do not add prefix like ['Respone', 'based on the document']. Do not add any further Explanation,Note.) '''
    
                    
                else:
                    prompt_1 = f'''Act as a financial analyst and give concise answer to below Question as truthfully as possible, with given Context.
                                Do not provide any extra [Explanation, Note,Description] block below the Response.\n\n\
                                Question: {query}\n\
                                Context: {context_1}\n\                      
                                Response: (Act like a professional and provide me a concise Response . Do not add any extra [Explanation, Note, Descricption] below the Response.)'''
    
    
                #prompt = PromptTemplate(template=prompt, input_variables=["query", "context"])
                # response = usellm(prompt_1) #LLM_Response()
                response = llama_llm(llama_13b,prompt_1)
                text_dict[query] = response
    
                st.write(response)
    
                if response:
                    df = pd.DataFrame(text_dict.items(), columns=['Question','Answer'])
                else:
                    df = pd.DataFrame()
    
                st.session_state["tmp_table_llama_fd"] = pd.concat([st.session_state.tmp_table_llama_fd, df], ignore_index=True)
                st.session_state.tmp_table_llama_fd.drop_duplicates(subset=['Question'])
    
   
    
    # with st.spinner("Downloading...."):
    # if st.button("Download Response", disabled=st.session_state.disabled):
    # Create a Word document with the table and some text
    
    # col_d1, col_d2 = st.columns(2)
    col_s1, col_s2, col_d1, col_d2 = st.tabs(["Summarize","SAR Narrative","Download Report", "Download Case Package"])
    
    with col_s1:
        with st.spinner('Summarization ...'):
            if st.button("Summarize",disabled=st.session_state.disabled):
                if st.session_state.llm == "Open-AI":
                    st.session_state.disabled=False
            
                    summ_dict_gpt = st.session_state.tmp_table_gpt_fd.set_index('Question')['Answer'].to_dict()
                    # chat_history = resp_dict_obj['Summary']
                    memory = ConversationSummaryBufferMemory(llm=llm, max_token_limit=300)
                    memory.save_context({"input": "This is the entire summary"}, 
                                    {"output": f"{summ_dict_gpt}"})
                    conversation = ConversationChain(
                    llm=llm, 
                    memory = memory,
                    verbose=True)
                    st.session_state["tmp_summary_gpt_fd"] = conversation.predict(input="Provide a detailed summary of the text provided by reframing the sentences. Provide the summary in a single paragraph. Please don't include words like these: 'chat summary', 'includes information' in my final summary.")
                    # showing the text in a textbox
                    # usr_review = st.text_area("", value=st.session_state["tmp_summary_gpt"])
                    # if st.button("Update Summary"):
                    #     st.session_state["fin_opt"] = usr_review
                    st.write(st.session_state["tmp_summary_gpt_fd"])
    
    
                elif st.session_state.llm == "Open-Source":
                    st.session_state.disabled=False
                    template = """Write a detailed summary.
                    Return your response in a single paragraph.
                    ```{text}```
                    Response: """
                    prompt = PromptTemplate(template=template,input_variables=["text"])
                    llm_chain_llama = LLMChain(prompt=prompt,llm=llama_13b)
                    st.table(st.session_state.tmp_table_llama_fd)
                    summ_dict_llama = st.session_state.tmp_table_llama_fd.set_index('Question')['Answer']
                    text = []
                    for key,value in summ_dict_llama.items():
                        text.append(value)
                    st.session_state["tmp_summary_llama_fd"] = llm_chain_llama.run(text)
                    st.write(st.session_state["tmp_summary_llama_fd"])
    
        
        tmp_summary = []
        tmp_table = pd.DataFrame()
     
        if st.session_state.llm == "Open-AI":
            st.session_state.disabled=False
            tmp_table = pd.concat([tmp_table, st.session_state["tmp_table_gpt_fd"]], ignore_index=True)
            tmp_summary.append(st.session_state["tmp_summary_gpt_fd"])
    
    
        elif st.session_state.llm == "Open-Source":
            st.session_state.disabled=False
            tmp_summary.append(st.session_state["tmp_summary_llama_fd"])
            tmp_table = pd.concat([tmp_table, st.session_state["tmp_table_llama_fd"]], ignore_index=True)
    
    
        try:
            # initiate the doc file
            doc = docx.Document()
            # doc.add_section(WD_SECTION.NEW_PAGE)
            doc.add_heading(f"Case No.: {st.session_state.case_num}",0)
    
            # Add a subheader for case details
            subheader_case = doc.add_paragraph("Case Details")
            subheader_case.style = "Heading 2"
            # Addition of case details
            paragraph = doc.add_paragraph(" ")
            case_info = {
                "Case Number                            ": " SAR-2023-24680",
                "Customer Name                       ": " John Brown",
                "Customer ID                              ": " 9659754",
                "Case open date                         ": " Feb 02, 2021",
                "Case Type                                  ": " Fraud Transaction",
                "Case Status                                ": " Open"
            }
            for key_c, value_c in case_info.items():
                doc.add_paragraph(f"{key_c}: {value_c}")
            paragraph = doc.add_paragraph(" ")
    
            # Add a subheader for customer info to the document ->>
            subheader_paragraph = doc.add_paragraph("Customer Information")
            subheader_paragraph.style = "Heading 2"
            paragraph = doc.add_paragraph(" ")
    
            # Add the customer information
            customer_info = {
                "Name                                           ": " John Brown",
                "Address                                      ": " 858 3rd Ave, Chula Vista, California, 91911 US",
                "Phone                                          ": " (619) 425-2972",
                "A/C No.                                        ": " 4587236908230087",
                "SSN                                               ": " 653-30-9562"
            }
    
            for key, value in customer_info.items():
                doc.add_paragraph(f"{key}: {value}")
            paragraph = doc.add_paragraph()
            # Add a subheader for Suspect infor to the document ->>
            subheader_paragraph = doc.add_paragraph("Suspect's Info")
            subheader_paragraph.style = "Heading 2"
            paragraph = doc.add_paragraph()
            #""" Addition of a checkbox where unticked box imply unavailability of suspect info"""
    
            # Add the customer information
            sent_val = "Suspect has been reported."
            paragraph = doc.add_paragraph()
            runner = paragraph.add_run(sent_val)
            runner.bold = True
            runner.italic = True
            suspect_info = {
                "Name                                             ": "Mike White",
                "Address                                        ": "520, WintergreenCt,Vancaville,CA,95587",
                "Phone                                             ": "NA",
                "SSN                                                 ": "NA",
                "Relationship with Customer  ": "NA"
            }
    
            for key, value in suspect_info.items():
                doc.add_paragraph(f"{key}: {value}")
    
            doc.add_heading('Summary', level=2)
            paragraph = doc.add_paragraph()
            doc.add_paragraph(tmp_summary)
            paragraph = doc.add_paragraph()
            doc.add_heading('Key Insights', level=2)
            paragraph = doc.add_paragraph()
            tmp_table.drop_duplicates(inplace=True)
            columns = list(tmp_table.columns)
            table = doc.add_table(rows=1, cols=len(columns), style="Table Grid")
            table.autofit = True
            for col in range(len(columns)):
                # set_cell_margins(table.cell(0, col), top=100, start=100, bottom=100, end=50) # set cell margin
                table.cell(0, col).text = columns[col]
            # doc.add_table(st.session_state.tmp_table.shape[0]+1, st.session_state.tmp_table.shape[1], style='Table Grid')
    
            for i, row in enumerate(tmp_table.itertuples()):
                table_row = table.add_row().cells # add new row to table
                for col in range(len(columns)): # iterate over each column in row and add text
                    table_row[col].text = str(row[col+1]) # avoid index by adding col+1
            # save document
            # output_bytes = docx.Document.save(doc, 'output.docx')
            # st.download_button(label='Download Report', data=output_bytes, file_name='evidence.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
            bio = io.BytesIO()
            doc.save(bio)
        except NameError:
            pass
    

    with col_s2:
        with st.spinner('SAR Narrative ...'):
            if st.button("SAR Narrative",disabled=st.session_state.disabled):
                if st.session_state.llm == "Open-AI":
                    st.session_state.disabled=False
            
                    narrative_dict_gpt = st.session_state.tmp_narr_table_gpt.set_index('Question')['Answer'].to_dict()
                    memory = ConversationSummaryBufferMemory(llm=llm, max_token_limit=300)
                    memory.save_context({"input": "This is the entire summary"}, 
                                    {"output": f"{narrative_dict_gpt}"})
                    conversation = ConversationChain(
                    llm=llm, 
                    memory = memory,
                    verbose=True)
                  
                    st.session_state["tmp_narrative_gpt"] = conversation.predict(input=" You need to create a SAR Narrative which should include the answers to all the questions mentioned below in ():\
                                                                                          (1.What is the suspect's name ? \
                                                                                          2. When did the fraud occur? \
                                                                                          3.What type of fraud is taking place? \
                                                                                          4.What type of transaction is involved? \
                                                                                          5.How did the fraud occur? \
                                                                                          6.Is it a SAR case? If yes, then why?) \
                                                                                          Create a concise SAR narrative in not more than 100 words considering all the factors mentioned above. \
                                                                                          You can use the chat_history to answer these questions. Please don't include words like these: 'chat summary', 'includes information' in my SAR Narrative.")
                    # showing the text in a textbox
                    # usr_review = st.text_area("", value=st.session_state["tmp_summary_gpt"])
                    # if st.button("Update Summary"):
                    #     st.session_state["fin_opt"] = usr_review
                    st.write(st.session_state["tmp_narrative_gpt"])
    
    
                elif st.session_state.llm == "Open-Source":
                    st.session_state.disabled=False
                    template = """You need to create a SAR Narrative which should include the answers to all the questions mentioned below in ():\
                                                                                          (1.What is the suspect's name ? \
                                                                                          2. When did the fraud occur? \
                                                                                          3.What type of fraud is taking place? \
                                                                                          4.What type of transaction is involved? \
                                                                                          5.How did the fraud occur? \
                                                                                          6.Is it a SAR case? If yes, then why?) \
                                                                                          Create a concise SAR narrative in not more than 100 words considering all the factors mentioned above. \
                                                                                          You can use the chat_history to answer these questions. Please don't include words like these: 'chat summary', 'includes information' in my SAR Narrative.")
                    ```{text}```
                    Response: """
                    prompt = PromptTemplate(template=template,input_variables=["text"])
                    llm_chain_llama = LLMChain(prompt=prompt,llm=llama_13b)

                    narrative_dict_llama = st.session_state.tmp_narr_table_llama.set_index('Question')['Answer']
                    text = []
                    for key,value in narrative_dict_llama.items():
                        text.append(value)
                    st.session_state["tmp_narrative_llama"] = llm_chain_llama.run(text)
                    st.write(st.session_state["tmp_narrative_llama"]) 
    
        
        
    with col_d1:
    # Applying to download button -> download_button
        st.markdown("""
            <style>
                .stButton download_button {
                    width: 100%;
                    height: 70%;
                }
            </style>
        """, unsafe_allow_html=True)
    
        
        if doc:
            st.download_button(
                label="Download Report",
                data=bio.getvalue(),
                file_name="Report.docx",
                mime="docx",
                disabled=st.session_state.disabled
            )
    with col_d2:
        
        # initiating a temp file
        tmp_dir = tempfile.mkdtemp()
    
        file_paths= []
    
        for uploaded_file in st.session_state.pdf_files:
            file_pth = os.path.join(tmp_dir, uploaded_file.name)
            with open(file_pth, "wb") as file_opn:
                file_opn.write(uploaded_file.getbuffer())
            file_paths.append(file_pth)
    
        for fetched_pdf in fetched_files:
            # st.write(fetched_pdf)
            file_pth = os.path.join('data/', fetched_pdf)
            # st.write(file_pth)
            file_paths.append(file_pth)
    
        
        combined_doc_path = os.path.join(tmp_dir, "report.docx")
        doc.save(combined_doc_path)
    
    
    
        # Create a zip file with the uploaded PDF files and the combined document
        zip_file_name = "package_files.zip"
        if file_paths:
            files =  [combined_doc_path] + file_paths
            create_zip_file(files, zip_file_name)
            # create_zip_file(file_paths, zip_file_name)
        else:
            pass
    
        
        # Download the package
    
        with open(zip_file_name, "rb") as file:
            st.download_button(
                label="Download Case Package", 
                data=file, 
                file_name=zip_file_name,
                disabled=st.session_state.disabled)
    
            # # Cleanup: Delete the temporary directory and its contents
            # for file_path in file_paths + [combined_doc_path]:
            #     os.remove(file_path)
            # os.rmdir(temp_dir)
    
            
    # Adding Radio button
    st.header("Make Decision")
    st.markdown(
            """ <style>
                    div[role="radiogroup"] >  :first-child{
                        display: none !important;
                    }
                </style>
                """,
            unsafe_allow_html=True
        )
    st.markdown("##### Is SAR filing required?")
    selected_rad = st.radio(":blue[Please select your choice]", ["opt1","Yes", "No", "Refer for review"], horizontal=True,disabled=st.session_state.disabled)
    if selected_rad == "Refer for review":
        email_regex = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        email_id = st.text_input("Enter your email ID")
        if email_id and not re.match(email_regex, email_id):
            st.error("Please enter a valid email ID")
    if st.button("Submit"):
        if selected_rad in ("str_opt1"):
            st.write("")
        elif selected_rad in ("Yes"):
            st.warning("Thanks for your review, your response has been submitted")
        elif selected_rad in ("No"):
            st.success("Thanks for your review, your response has been submitted")
    
        else:
            st.info("Thanks for your review, Case has been assigned to the next reviewer")
    
    
    # # Allow the user to clear all stored conversation sessions
    # if st.button("Reset Session"):
    #     reset_session_state()
    #     # st.cache_data.clear()
    #     # pdf_files.clear()
    
    
    
    # Footer
    st.markdown(
        """
        <style>
          #MainMenu {visibility: hidden;}
          footer {visibility: hidden;}
        </style>
        """
        , unsafe_allow_html=True)
    st.markdown('<div class="footer"><p></p></div>', unsafe_allow_html=True)
    
    hide_st_style = """
                <style>
                #MainMenu {visibility: hidden;}
                footer {visibility: hidden;}
                header {visibility: hidden;}
                </style>
                """
    st.markdown(hide_st_style, unsafe_allow_html=True)
    
    
    padding = 0
    st.markdown(f""" <style>
        .reportview-container .main .block-container{{
            padding-top: {padding}rem;
            padding-right: {padding}rem;
            padding-left: {padding}rem;
            padding-bottom: {padding}rem;
        }} </style> """, unsafe_allow_html=True)
    
elif selected_option_case_type == "AML":
    
    st.markdown("### :red[Anti-Money Laundering]")
    if selected_option == "SAR-2023-24680":
        st.session_state.case_num = "SAR-2023-24680"
        
    
        col1,col2 = st.columns(2)
        # Row 1
        with col1:
            st.markdown("**Case number&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** SAR-2023-24680")
            st.markdown("**Customer name  :** Sarah Jones")
    
    
        with col2:
            st.markdown("**Case open date&nbsp;&nbsp;&nbsp;&nbsp;:** Sep 01, 2022")
            st.markdown("**Case type&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** Money Laundering")
    
    
        # Row 2
        with col1:
            st.markdown("**Customer ID&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** 560062")
    
    
        with col2:
            st.markdown("**Case Status&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** Open")
    
    
        # Evidence uploader/Fetch    
        st.header("Upload Evidence")
        
        directoty_path = "ml_doc/"
        fetched_files = read_pdf_files(directoty_path)
        if selected_option:
            # Create two columns

            
            col1_up, col2_up = st.tabs(["Fetch Evidence", "Upload Evidence"])
            with col1_up:
                
                if 'clicked' not in st.session_state:
                    st.session_state.clicked = False
                
                def set_clicked():
                    st.session_state.clicked = True
                    st.session_state.disabled = True
                
                st.button('Fetch Evidence', on_click=set_clicked)
    
                if st.session_state.clicked:
            
                    #select box to select file
                    selected_file_name = st.selectbox(":blue[Select a file to View]",fetched_files)
                    st.write("Selected File: ", selected_file_name)
                    st.session_state.disabled = False
                    file_ext = tuple("pdf")  
                    file_ext1 = tuple("xlsx")
                  
                    if selected_file_name.endswith(file_ext):
                        selected_file_path = os.path.join(directoty_path, selected_file_name)
                        #converting pdf data to bytes so that render_pdf_as_images could read it
                        file = pdf_to_bytes(selected_file_path)
                        pdf_images = render_pdf_as_images(file)
                        #showing content of the pdf
                        st.subheader(f"Contents of {selected_file_name}")
                        for img_bytes in pdf_images:
                            st.image(img_bytes, use_column_width=True)
                        
                    elif selected_file_name.endswith(file_ext1):
                        selected_file_path = os.path.join(directoty_path, selected_file_name)
                        # from openpyxl import reader,load_workbook,Workbook
                        # wb=Workbook()
                        # wb = load_workbook(selected_file_path, read_only=False)
                        # st.write(wb.sheetnames)
                        # st.title(wb)
                        # st.write(wb.active)
                        # Read the Excel file into a DataFrame
                        df = pd.read_excel(selected_file_path, engine='openpyxl')
                        # Find the row index where the table data starts
                        data_start_row = 0  # Initialize to 0
                        for i, row in df.iterrows():
                            if row.notna().all():
                                data_start_row = i
                                break       
                        if data_start_row>0:  
                          df.columns = df.iloc[data_start_row]
                        # Extract the text content above the data
                        text_content = "\n".join(df.iloc[:data_start_row].apply(lambda x: "\t".join(map(str, x)), axis=1)).replace('nan','')
                        
                        
                        st.text(text_content)
                        
                        st.write(df.iloc[data_start_row+1:], index=False)
                
                 
                    else:
                        selected_file_path = os.path.join(directoty_path, selected_file_name)
                        # This is showing png,jpeg files
                        st.image(selected_file_path, use_column_width=True)
    
    
    
            with col2_up:
                pdf_files = st.file_uploader("", type=["pdf","png","jpeg","docx","xlsx"], accept_multiple_files=True)
                st.session_state.pdf_files = pdf_files
                # showing files
                for uploaded_file in pdf_files:
                    #This code is to show pdf files
                    file_ext = tuple("pdf")
                    if uploaded_file.name.endswith(file_ext):
                        # Show uploaded files in a dropdown
                        # if pdf_files:
                        st.subheader("Uploaded Files")
                        file_names = [file.name for file in pdf_files]
                        selected_file = st.selectbox(":blue[Select a file]", file_names)
                        # Enabling the button
                        st.session_state.disabled = False
                        # Display selected PDF contents
                        if selected_file:
                            selected_pdf = [pdf for pdf in pdf_files if pdf.name == selected_file][0]
                            pdf_images = render_pdf_as_images(selected_pdf)
                            st.subheader(f"Contents of {selected_file}")
                            for img_bytes in pdf_images:
                                st.image(img_bytes, use_column_width=True)
    
                    else:
                        # This is showing png,jpeg files
                        st.image(uploaded_file, use_column_width=True)
    
    #creating temp directory to have all the files at one place for accessing
        tmp_dir_ = tempfile.mkdtemp()
        temp_file_path= []
    
        for uploaded_file in pdf_files:
            file_ext = tuple("pdf")
            if uploaded_file.name.endswith(file_ext):
                file_pth = os.path.join(tmp_dir_, uploaded_file.name)
                with open(file_pth, "wb") as file_opn:
                    file_opn.write(uploaded_file.getbuffer())
                    temp_file_path.append(file_pth)
            else:
                pass
    
    
        for fetched_pdf in fetched_files:
            file_ext = ("pdf","xlsx")
            if fetched_pdf.endswith(file_ext):
                file_pth = os.path.join('ml_doc/', fetched_pdf)
                # st.write(file_pth)
                temp_file_path.append(file_pth) 
            else:
                pass
    
        #combining files in fetch evidence and upload evidence
        pdf_files_ = []
        if temp_file_path:
            if pdf_files and fetched_files:
                file_names = [file.name for file in pdf_files]
                file_names = file_names + fetched_files
                pdf_files_ = file_names
            elif fetched_files:
                pdf_files_ = fetched_files
            elif pdf_files:
                pdf_files_ = pdf_files
            else: pass
    
         #This is the embedding model
        model_name= "thenlper/gte-small"
        
        
        # Memory setup for gpt-3.5
        llm = ChatOpenAI(temperature=0.1)
        memory = ConversationSummaryBufferMemory(llm=llm, max_token_limit=500)
        conversation = ConversationChain(llm=llm, memory =memory,verbose=False)
        
        
        # Adding condition on embedding
        try:
            if temp_file_path:
                hf_embeddings = embed(model_name) 
            else:
                pass
        except NameError:
            pass
        
        # Chunking with overlap
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size = 1000,
            chunk_overlap  = 100,
            length_function = len,
            separators=["\n\n", "\n", " ", ""]
        )
        
    
    else:
        # Disabling the button
        st.session_state.disabled = True
        st.session_state.case_num = selected_option    
    
    # Creating header
    col1,col2 = st.columns(2)
    with col1:
        st.subheader('Pre-Set Questionnaire')
        # Create a Pandas DataFrame with your data
    
        data = {'Questions': ["Is there is any suspicious activity?",
                              "If there is any suspicious activity taking place, What is the Suspect Name?",
          # "Is there any potential Money Laundering activity based on the transaction statements",
                          "What are the transaction that can be associated with Money Laundering activity?",
                          "When is the Money laundering activity taking place?",
                         "What type of Money laundering activity is taking place?",
                          "What is the total amount associated with the money laundering activity?"]}
        df_fixed = pd.DataFrame(data)
        df_fixed.index = df_fixed.index +1
    with col2:
        # Create a checkbox to show/hide the table
        cols1, cols2, cols3, cols4 = st.columns([1,1,1,1])
        with cols1:
            show_table = tog.st_toggle_switch(label="", 
                                key="Key1", 
                                default_value=False, 
                                label_after = False, 
                                inactive_color = '#D3D3D3', 
                                active_color="#11567f", 
                                track_color="#29B5E8"
                                )
        # Show the table if the checkbox is ticked
        if show_table:
            # st.write(df_fixed)
            # st.dataframe(df_fixed, width=1000)
            df_fixed["S.No."] = df_fixed.index
            df_fixed = df_fixed.loc[:,['S.No.','Questions']]
            st.markdown(df_fixed.style.hide(axis="index").to_html(), unsafe_allow_html=True)
    
    with st.spinner('Wait for it...'):
        if st.button("Generate Insights",disabled=st.session_state.disabled):
            if temp_file_path is not None:
                # File handling logic
                _, docsearch = embedding_store(temp_file_path)
                if st.session_state.llm == "Open-AI":
                    chat_history_1 = {}

                    query = "Is there any Suspicious Activity?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 = f'''You are a financial analyst, who is an expert in detecting Fraud and money-laundering activities go through the following savings account as well as credit card transactions statement \
                    and detect if any Suspicious activity is taken place or not. Suspicious activity can be of the type Fraud or Money-Laundering. Any transaction amount in the statement above \
                    or equal to $10,000 can be a indication of  Money-Laundering activity.\
                    Select the most appropriate Type while giving the response.\n\n\
                            Question: {query}\n\
                            Context: {context_1}\n\
                            Response: '''
                    response = usellm(prompt_1)
                    chat_history_1[query] = response

                    query = "If there is any suspicious activity taking place, What is the Suspect Name?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 = f'''Perform Name Entity Recognition to find out the Suspect Name, given the savings account and credit \
                    card transaction statement.
                    Customer name can be taken as the Suspect name. \
                    Give only the name entity and nothing else.\n\n\
                            Question: {query}\n\
                            Context: {context_1}\n\
                            Response: '''

                    # prompt_1=f'''You Are an Anti-Money Laundering Specialist, go through the credit card and savings account transaction statement and look for a suspicious \
                    # activity. Any transaction above or equal to $10,000 in the statement can be considered to be suspicious activity. \
                    # If there is any suspicious transaction present, give your response as yes, along with the customer name which can be considered as suspect name.\
                    # Question: {query}\n\
                    # Context: {context_1}\n\
                    # Response: '''
                  
                    response = usellm(prompt_1)
                    chat_history_1[query] = response
                  
                    # query = "Is there any potential Money Laundering activity based on the transaction statements?"
                    # context_1 = docsearch.similarity_search(query, k=5)
                    # prompt_1 = f'''You Are an Anti-Money Laundering Specialist who is an expert in detecting Money-laundering. \n
                    # You need to look closely into the credit card transaction statements as well as savings account transaction statements collectively and evaluate \
                    # them together to check for any potential suspicious money laundering activities. \n
                    # A Money laundering activity can be detected if any of the following transaction patterns is observed-:
                    # 1) If there are cash transactions happening, greater than or equal to $10,000.
                    # 2) If there is a high-value international transaction happening which involves movement of funds to or from a high risk geographical location(Ex- Mauritious, Syria, Nigeria,etc.).
                    # 3) If there is any money laundering pattern like structuring or smurfing, layering, placement, integration, etc observed within 
                    # the credit card and savings bank account transactions statements collectively.
                    # Provide your response as Yes if there is a hint of Money being Laundered considering all of the factors above.\n\n\
                    #         Question: {query}\n\
                    #         Context: {context_1}\n\
                    #         Response: '''
                    # response = usellm(prompt_1)
                    # chat_history_1[query] = response
    
    
                    query = "What are the transaction that can be associated with Money Laundering activity?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f'''You Are an Anti-Money Laundering Specialist, Identify the transactions \
                                that can be potentially associated with the Money Laundering activity both from Credit Card transaction statement as well as savings account statement collectively. \n
                                Money laundering transactions often involve characteristics like large cash deposits greater than or equal to $10,000 \
                                Payments greater than or equal to 10000$ to an unrecognized entity with no specific business purpose, \ 
                                , transactions involving movement of funds to or from high-risk locations(Ex- Mauritious, Syria, Nigeria,etc.) and are greater than 10000$, any suspicion of money laundered via structuring , layering or intergration, process, \
                                Cash deposits greater than or equal to 10000$ with source of funds not clear used to pay off credit card debt, etc\n \n
                                Only include transactions which are greater than or equal to 10,000$ in your response. \n
                                Give all such suspicious transactions grouped by transaction type(Credit card, savings account,etc.) along with dates, amounts from the context as your response \
                                Do not include any credit card payment transaction in your response. Do not repeat the above information and provide a to the point response.\n\n
                                Context: {context_1}\n\
                                Response: (Give me a concise response .Do not give me any Explanation,Note, etc.)'''

                    response = usellm(prompt_1)
                    chat_history_1[query] = response

                    query = "When is the Money laundering activity taking place?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    # prompt_1 =  f'''You Are an Anti-Money Laundering Specialist, Identify all the dates on which transactions \
                    #             that can be potentially associated with the Money Laundering activity both from Credit Card transaction statement as well as savings account statement collectively is happening. \n
                    #             Money laundering transactions often involve characteristics like large cash deposits greater than or equal to $10,000 \
                    #             Payments greater than or equal to 10000$ to an unrecognized entity with no specific  business purpose, \ 
                    #             ,Credit card transactions involving movement of funds to or from high-risk locations(Ex- Mauritious, Syria, Nigeria,etc.), any suspicion of money laundered via structuring , layering or intergration, process, \
                    #             Cash deposits with source of funds not clear used to pay off debt, etc. \n
                    #             Do not consider transactions less than 10000$ for creating the response. \n
                    #             Give the dates only of all such suspicious transactions grouped by transaction type(Credit card Transaction, savings account transaction,etc.) from the context as your response \
                    #             Do not include any credit card payment transaction present in the savings transaction statement in your response. Also, Do not repeat the above information and provide a to the point response. Also, do not include transactions less than 10000$ in your response.\n\n
                    #             Context: {context_1}\n\
                    #             Response: (Give me a concise response .Do not give me any Explanation,Note, etc.)'''

                    prompt_1= f'''You Are an Anti-Money Laundering Specialist, Give all the dates on which the transactions that can be potentially associated with \
                                  the Money laundering activity from the credit card transactions statement or the savings account transaction statement. Money Laundering activity often involves characteristics like \
                                  large cash deposits or transfer greater or equal to $10,000 to an unrecognized entity with no specific business purpose, transactions \
                                  involving movement of funds to or from high-risk locations(Ex- Mauritious, Syria, Nigeria,etc.). Give the dates only of all such suspicious transactions.\
                                  Context: {context_1}\n\
                                  Response: (Give me a concise response .Do not give me any Explanation,Note, etc.)'''

                    response = usellm(prompt_1)
                    chat_history_1[query] = response

                    query = "What type of Money laundering activity is taking place?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f'''You Are an Anti-Money Laundering Specialist, give the \
                                type of money laundering activity that is taking place based on the transaction \
                                patterns observed in credit card and savings account transaction statements. The type may include Layering, Structuring, Round-tripping etc. \
                                Look carefully into the transactions statements mentioned above and give a precise answer with explanation of why you think a specific type of money laundering is happening..\n\n\
                                Context: {context_1}\n\
                                Response: (Give me a concise response in not more than 50 words.Do not give me any Explanation,Note)'''

                    response = usellm(prompt_1)
                    chat_history_1[query] = response

                    query = "What is the total amount associated with the money laundering activity?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f'''You Are an Anti-Money Laundering Specialist, Identify the transactions \
                                that can be potentially associated with the Money Laundering activity both from Credit Card transaction statement as well as savings account statements collectively. \n
                                Money laundering transactions often involve characteristics like large cash deposits greater than or equal to $10,000, \
                                Payments greater than or equal to 10000$ to an unrecognized entity with no specific  business purpose, \ 
                                , transactions involving movement of funds to or from high-risk locations(Ex- Mauritious, Syria, Nigeria,etc.) and are greater than 10000$, any suspicion of money laundered via structuring , layering or intergration, process, \
                                Cash deposits greater than or equal to 10000$ with source of funds not clear used to pay off credit card debt, etc\n \n
                                Only include transactions which are greater than or equal to 10,000$ in your response. \n
                                Add the dollar amount associated with all such suspicious transactions to get the total amount associated \
                                with the money laundering activity.
                                Only include the total amount in your response .\n\n\
                                Context: {context_1}\n\
                                Response: (Give me a concise response in one sentence.Do not give me any Explanation,Note)'''
                    
                    response = usellm(prompt_1)
                    chat_history_1[query] = response

    
                    try:
                        res_df_gpt = pd.DataFrame(list(chat_history_1.items()), columns=['Question','Answer'])
                        res_df_gpt.reset_index(drop=True, inplace=True)
                        index_ = pd.Series([1,2,3,4,5,6])
                        res_df_gpt = res_df_gpt.set_index([index_])
                        # st.write(res_df_gpt)
                    except IndexError: 
                        pass
                    st.table(res_df_gpt)
                    st.session_state["tmp_table_gpt_aml"] = pd.concat([st.session_state.tmp_table_gpt_aml, res_df_gpt], ignore_index=True)
                
                
                elif st.session_state.llm == "Open-Source":
    
                    chat_history = {}
    
                    query = "Is there any potential Money Laundering activity based on the transaction statements?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 = f'''You Are an Anti-Money Laundering Specialist who is an expert in detecting Money-laundering. 
                    A Money laundering activity can be detected if any of the following transaction patterns is observed-:
                    1) If there are multiple transactions happening, greater than or equal to $10,000 in a short span of time.
                    2) If there is a high-value international transaction happening which involves a high risk geographical location.
                    3) If there is any money laundering pattern like structuring or smurfing, layering, placement, integration, etc observed within 
                    the transactions statement.
                    Provide your response as Yes if there is a hint of Money being Laundered considering all of the factors above.\n\n\
                            Question: {query}\n\
                            Context: {context_1}\n\
                            Response: (Give me response in one sentence. Do not give me any Explanation or Note)'''
                  
                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response
                  
    
                    query = "What are the transaction that can be associated with Money Laundering activity?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f'''You Are an Anti-Money Laundering Specialist, Identify the transactions \
                                that can be potentially associated with the Money Laundering activity. Money laundering \
                                transactions often involve characteristics like large cash deposits, High value transactions greater than or equal to $10,000 \
                                within a short span of time, transactions with high-risk countries, money laundered via structuring process. Give precise response, \
                                do not include any other unnecessary information or the Balance amount.\n\n
                
                                Context: {context_1}\n\
                                Response: (Give me a concise response in one sentence.Do not give me any Explanation,Note)'''

                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response

                    query = "When is the Money laundering activity taking place?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f'''You Are an Anti-Money Laundering Specialist, give all the dates when a money laundering activity is taking place given the context. Money laundering transactions often \
                                involve characteristics like large cash deposits in savings tansaction summary equal and above $10,000, \
                                rapid movement of funds, transactions with high-risk countries, or unexplained source of funds. Specifically, all transactions above or \ 
                                equal to $10,000 are considered to be a potential money laundering transaction. Answer the question considering the factors mentioned above with transaction details.\n\n\
                                Context: {context_1}\n\
                                Response: (Give me a concise response in one sentence.Do not give me any Explanation,Note)'''

                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response
                  
                    query = "What type of Money laundering activity is taking place?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f'''You Are an Anti-Money Laundering Specialist, give the \
                                type of money laundering activity that is taking place based on the transaction \
                                patterns observed. The type may include Layering, Structuring, Round-tripping etc. \
                                Look carefully into the transactions statement and give a precise answer with explanation of why you think a specific type of money laundering is happening..\n\n\
                                Context: {context_1}\n\
                                Response: (Give me a concise response in not more than 50 words.Do not give me any Explanation,Note)'''

                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response

                    query = "What is the total amount associated with the money laundering activity?"
                    context_1 = docsearch.similarity_search(query, k=5)
                    prompt_1 =  f'''You Are an Anti-Money Laundering Specialist, give the total amount \
                                associated with money laundering activity that is taking place Based on the \
                                transaction statement, for getting the total amount, you can add all the money laundering \
                                transactions amount.\n\n\
                                Context: {context_1}\n\
                                Response: (Give me a concise response in one sentence.Do not give me any Explanation,Note)'''
                    
                    response = llama_llm(llama_13b,prompt_1)
                    chat_history[query] = response
                
                   

                    try:
                        res_df_llama = pd.DataFrame(list(chat_history.items()), columns=['Question','Answer'])
                        res_df_llama.reset_index(drop=True, inplace=True)
                        index_ = pd.Series([1,2,3,4,5])
                        res_df_llama = res_df_llama.set_index([index_])
                        # st.write(res_df_llama)
                  
                    except IndexError: 
                        pass
                    st.table(res_df_llama)
                    st.session_state["tmp_table_llama_aml"] = pd.concat([st.session_state.tmp_table_llama_aml, res_df_llama], ignore_index=True)
                
                
    
    
    st.markdown("---")
    
    # For input box outside of template4
    try:
        if temp_file_path:
            docs, docsearch = embedding_store(temp_file_path)
        else:
            pass
    except Exception:
        pass
    
    
    # Text Input
    st.subheader("Ask Additional Questions")
    query = st.text_input(':blue[Please ask below the additional case questions.]',disabled=st.session_state.disabled)
    text_dict = {}
    @st.cache_data
    def LLM_Response():
        llm_chain = LLMChain(prompt=prompt, llm=llm)
        response = llm_chain.run({"query":query, "context":context})
        return response
    if st.session_state.llm == "Open-AI":
        with st.spinner('Getting you information...'):      
            if query:
                # Text input handling logic
                #st.write("Text Input:")
                #st.write(text_input)
    
                context_1 = docsearch.similarity_search(query, k=5)
                st.session_state.context_1 = context_1
                
            
                prompt_1 = f''' You Are an Anti-Money Laundering Specialist, provide the answer to the below question in a concise manner.\n\n\
                            Question: {query}\n\
                            Context: {context_1}\n\                      
                            Response: '''
    
    
                #prompt = PromptTemplate(template=prompt, input_variables=["query", "context"])
                response = usellm(prompt_1) #LLM_Response()
                text_dict[query] = response
                # resp_dict_obj.update(text_dict)
                st.write(response)
                if response:
                    df = pd.DataFrame(text_dict.items(), columns=['Question','Answer'])
                else:
                    df = pd.DataFrame()
    
                st.session_state["tmp_table_gpt_aml"] = pd.concat([st.session_state.tmp_table_gpt_aml, df], ignore_index=True)
                st.session_state.tmp_table_gpt_aml.drop_duplicates(subset=['Question'])
    
    
    elif st.session_state.llm == "Open-Source":
        with st.spinner('Getting you information...'):      
            if query:
                # Text input handling logic
                #st.write("Text Input:")
                #st.write(text_input)
    
                context_1 = docsearch.similarity_search(query, k=5)
                st.session_state.context_1 = context_1
                prompt_1 = f''' You Are an Anti-Money Laundering Specialist, provide the answer to the below question in a concise manner.\n\n\
                            Question: {query}\n\
                            Context: {context_1}\n\                      
                            Response: '''
    
                #prompt = PromptTemplate(template=prompt, input_variables=["query", "context"])
                # response = usellm(prompt_1) #LLM_Response()
                response = llama_llm(llama_13b,prompt_1)
                text_dict[query] = response
    
                st.write(response)
    
                if response:
                    df = pd.DataFrame(text_dict.items(), columns=['Question','Answer'])
                else:
                    df = pd.DataFrame()
    
                st.session_state["tmp_table_llama_aml"] = pd.concat([st.session_state.tmp_table_llama_aml, df], ignore_index=True)
                
                st.session_state.tmp_table_llama_aml.drop_duplicates(subset=['Question'])
    
    
    
    # col_d1, col_d2 = st.columns(2)
    col_s1, col_d1, col_d2 = st.tabs(["Summarize","Download Report", "Download Case Package"])
    
    with col_s1:
        with st.spinner('Summarization ...'):
            if st.button("Summarize",disabled=st.session_state.disabled):
                if st.session_state.llm == "Open-AI":
                    st.session_state.disabled=False
            
                    summ_dict_gpt = st.session_state.tmp_table_gpt_aml.set_index('Question')['Answer'].to_dict()
                    # chat_history = resp_dict_obj['Summary']
                    memory = ConversationSummaryBufferMemory(llm=llm, max_token_limit=300)
                    memory.save_context({"input": "This is the entire summary"}, 
                                    {"output": f"{summ_dict_gpt}"})
                    conversation = ConversationChain(
                    llm=llm, 
                    memory = memory,
                    verbose=True)
                    st.session_state["tmp_summary_gpt_aml"] = conversation.predict(input="Provide a detailed summary of the text provided by reframing the sentences. Provide the summary in a single paragraph. Please don't include words like these: 'chat summary', 'includes information' in my final summary.")
                    # showing the text in a textbox
                    # usr_review = st.text_area("", value=st.session_state["tmp_summary_gpt"])
                    # if st.button("Update Summary"):
                    #     st.session_state["fin_opt"] = usr_review
                    st.write(st.session_state["tmp_summary_gpt_aml"])
    
    
                elif st.session_state.llm == "Open-Source":
                    
                    st.session_state.disabled=False
                    template = """Write a detailed summary.
                    Return your response in a single paragraph.
                    ```{text}```
                    Response: """
                    prompt = PromptTemplate(template=template,input_variables=["text"])
                    llm_chain_llama = LLMChain(prompt=prompt,llm=llama_13b)
                    
                    summ_dict_llama = st.session_state.tmp_table_llama_aml.set_index('Question')['Answer']
                    text = []
                    for key,value in summ_dict_llama.items():
                        text.append(value)
                    st.session_state["tmp_summary_llama_aml"] = llm_chain_llama.run(text)
                    st.write(st.session_state["tmp_summary_llama_aml"])
    
        
        tmp_summary = []
        tmp_table = pd.DataFrame()
    
        if st.session_state.llm == "Open-AI":
            st.session_state.disabled=False
            tmp_table = pd.concat([tmp_table, st.session_state["tmp_table_gpt_aml"]], ignore_index=True)
            tmp_summary.append(st.session_state["tmp_summary_gpt_aml"])
    
    
        elif st.session_state.llm == "Open-Source":
            st.session_state.disabled=False
            tmp_summary.append(st.session_state["tmp_summary_llama_aml"])
            tmp_table = pd.concat([tmp_table, st.session_state["tmp_table_llama_aml"]], ignore_index=True)
    
    
        try:
            # initiate the doc file
            doc = docx.Document()
            # doc.add_section(WD_SECTION.NEW_PAGE)
            doc.add_heading(f"Case No.: {st.session_state.case_num}",0)
    
            # Add a subheader for case details
            subheader_case = doc.add_paragraph("Case Details")
            subheader_case.style = "Heading 2"
            # Addition of case details
            paragraph = doc.add_paragraph(" ")
            case_info = {
                "Case Number                            ": " SAR-2023-24680",
                "Customer Name                       ": " Sarah Jones",
                "Customer ID                              ": " 560062",
                "Case open date                         ": " Sep 01, 2022",
                "Case Type                                  ": " Money Laundering",
                "Case Status                                ": " Open"
            }
            for key_c, value_c in case_info.items():
                doc.add_paragraph(f"{key_c}: {value_c}")
            paragraph = doc.add_paragraph(" ")
    
            # Add a subheader for customer info to the document ->>
            subheader_paragraph = doc.add_paragraph("Customer Information")
            subheader_paragraph.style = "Heading 2"
            paragraph = doc.add_paragraph(" ")
    
            # Add the customer information
            customer_info = {
                "Name                                           ": " Sarah Jones",
                "Address                                      ": " 858 3rd Ave, Chula Vista, California, 91911 US",
                "Phone                                          ": " (619) 425-2972",
                "A/C No.                                        ": " 4587236908230087",
                "SSN                                               ": " 653-30-9562"
            }
    
            for key, value in customer_info.items():
                doc.add_paragraph(f"{key}: {value}")
            paragraph = doc.add_paragraph()
            # Add a subheader for Suspect infor to the document ->>
            subheader_paragraph = doc.add_paragraph("Suspect's Info")
            subheader_paragraph.style = "Heading 2"
            paragraph = doc.add_paragraph()
            #""" Addition of a checkbox where unticked box imply unavailability of suspect info"""
    
    
            doc.add_heading('Summary', level=2)
            paragraph = doc.add_paragraph()
            doc.add_paragraph(tmp_summary)
            paragraph = doc.add_paragraph()
            doc.add_heading('Key Insights', level=2)
            paragraph = doc.add_paragraph()
            tmp_table.drop_duplicates(inplace=True)
            columns = list(tmp_table.columns)
            table = doc.add_table(rows=1, cols=len(columns), style="Table Grid")
            table.autofit = True
            for col in range(len(columns)):
                # set_cell_margins(table.cell(0, col), top=100, start=100, bottom=100, end=50) # set cell margin
                table.cell(0, col).text = columns[col]
            # doc.add_table(st.session_state.tmp_table.shape[0]+1, st.session_state.tmp_table.shape[1], style='Table Grid')
    
            for i, row in enumerate(tmp_table.itertuples()):
                table_row = table.add_row().cells # add new row to table
                for col in range(len(columns)): # iterate over each column in row and add text
                    table_row[col].text = str(row[col+1]) # avoid index by adding col+1
            # save document
            # output_bytes = docx.Document.save(doc, 'output.docx')
            # st.download_button(label='Download Report', data=output_bytes, file_name='evidence.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
            bio = io.BytesIO()
            doc.save(bio)
        except NameError:
            pass
    
                
    with col_d1:
    # Applying to download button -> download_button
        st.markdown("""
            <style>
                .stButton download_button {
                    width: 100%;
                    height: 70%;
                }
            </style>
        """, unsafe_allow_html=True)
    
    
       
        
        if doc:
            st.download_button(
                label="Download Report",
                data=bio.getvalue(),
                file_name="Report.docx",
                mime="docx",
                disabled=st.session_state.disabled
            )
    with col_d2:
        
        # initiating a temp file
        tmp_dir = tempfile.mkdtemp()
    
        file_paths= []
    
        for uploaded_file in st.session_state.pdf_files:
            file_pth = os.path.join(tmp_dir, uploaded_file.name)
            with open(file_pth, "wb") as file_opn:
                file_opn.write(uploaded_file.getbuffer())
            file_paths.append(file_pth)
    
        for fetched_pdf in fetched_files:
            # st.write(fetched_pdf)
            file_pth = os.path.join('ml_doc/', fetched_pdf)
            # st.write(file_pth)
            file_paths.append(file_pth)
    
        
        combined_doc_path = os.path.join(tmp_dir, "report.docx")
        doc.save(combined_doc_path)
    
    
    
        # Create a zip file with the uploaded PDF files and the combined document
        zip_file_name = "package_files.zip"
        if file_paths:
            files =  [combined_doc_path] + file_paths
            create_zip_file(files, zip_file_name)
            # create_zip_file(file_paths, zip_file_name)
        else:
            pass
    
        
        # Download the package
    
        with open(zip_file_name, "rb") as file:
            st.download_button(
                label="Download Case Package", 
                data=file, 
                file_name=zip_file_name,
                disabled=st.session_state.disabled)
    
            # # Cleanup: Delete the temporary directory and its contents
            # for file_path in file_paths + [combined_doc_path]:
            #     os.remove(file_path)
            # os.rmdir(temp_dir)
    
            
    # Adding Radio button
    st.header("Make Decision")
    st.markdown(
            """ <style>
                    div[role="radiogroup"] >  :first-child{
                        display: none !important;
                    }
                </style>
                """,
            unsafe_allow_html=True
        )
    st.markdown("##### Is SAR filing required?")
    selected_rad = st.radio(":blue[Please select your choice]", ["opt1","Yes", "No", "Refer for review"], horizontal=True,disabled=st.session_state.disabled)
    if selected_rad == "Refer for review":
        email_regex = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        email_id = st.text_input("Enter your email ID")
        if email_id and not re.match(email_regex, email_id):
            st.error("Please enter a valid email ID")
    if st.button("Submit"):
        if selected_rad in ("str_opt1"):
            st.write("")
        elif selected_rad in ("Yes"):
            st.warning("Thanks for your review, your response has been submitted")
        elif selected_rad in ("No"):
            st.success("Thanks for your review, your response has been submitted")
    
        else:
            st.info("Thanks for your review, Case has been assigned to the next reviewer")
    
    
    # # Allow the user to clear all stored conversation sessions
    # if st.button("Reset Session"):
    #     reset_session_state()
    #     # st.cache_data.clear()
    #     # pdf_files.clear()
    
    
    
    # Footer
    st.markdown(
        """
        <style>
          #MainMenu {visibility: hidden;}
          footer {visibility: hidden;}
        </style>
        """
        , unsafe_allow_html=True)
    st.markdown('<div class="footer"><p></p></div>', unsafe_allow_html=True)
    
    hide_st_style = """
                <style>
                #MainMenu {visibility: hidden;}
                footer {visibility: hidden;}
                header {visibility: hidden;}
                </style>
                """
    st.markdown(hide_st_style, unsafe_allow_html=True)
    
    
    padding = 0
    st.markdown(f""" <style>
        .reportview-container .main .block-container{{
            padding-top: {padding}rem;
            padding-right: {padding}rem;
            padding-left: {padding}rem;
            padding-bottom: {padding}rem;
        }} </style> """, unsafe_allow_html=True)
st.markdown('---')



